"""
Pull Khmer text out of uploaded documents so it can be checked in the editor.

PDF and .docx carry their text as text, so extraction is exact and costs
nothing. Images are a different problem: the characters have to be recognised
first, and Khmer's stacked subscripts are exactly what OCR handles worst. The
endpoint accepts images so the frontend contract does not have to change when
that lands, but for now it says so plainly rather than returning bad text.

An extraction mistake and a spelling mistake look identical once the text is
in the editor, so it is better to return nothing than to return something
garbled that the spellchecker will then flag as the user's error.
"""

import base64
import io
import logging

from .khmer import has_khmer

log = logging.getLogger(__name__)

MAX_BYTES = 10 * 1024 * 1024  # 10 MB

PDF_TYPES = {"application/pdf"}
DOCX_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
TEXT_TYPES = {"text/plain", "text/markdown"}
IMAGE_TYPES = {"image/jpeg", "image/png", "image/webp", "image/heic", "image/gif"}


class ExtractionError(Exception):
    """Raised with a message intended to be shown to the user, in Khmer."""


def extract(data: bytes, mime_type: str, filename: str = "") -> tuple[str, str]:
    """
    Returns (text, note). The note is a user-facing warning, or empty.

    Dispatches on mime type, falling back to the filename extension because
    browsers report inconsistent types for .docx in particular.
    """
    if len(data) > MAX_BYTES:
        raise ExtractionError("ឯកសារធំពេក (អតិបរមា ១០ MB)")

    kind = _classify(mime_type, filename)

    if kind == "pdf":
        text, note = _from_pdf(data)
    elif kind == "docx":
        text, note = _from_docx(data), ""
    elif kind == "text":
        text, note = _from_text(data), ""
    elif kind == "image":
        from .ocr import OcrUnavailable, read_image

        try:
            text = read_image(data)
        except OcrUnavailable as exc:
            raise ExtractionError("ការអានអត្ថបទពីរូបភាពមិនអាចប្រើបានទេ") from exc

        if not text.strip():
            raise ExtractionError(
                "រកមិនឃើញអក្សរខ្មែរក្នុងរូបភាពនេះទេ។ "
                "សូមថតឱ្យច្បាស់ និងមានពន្លឺគ្រប់គ្រាន់។"
            )
        note = "អត្ថបទបានមកពីរូបភាព សូមពិនិត្យមុនប្រើ"
    else:
        raise ExtractionError(
            "ប្រភេទឯកសារនេះមិនត្រូវបានទ្រទ្រង់ទេ។ សូមប្រើ PDF, Word ឬ អត្ថបទ។"
        )

    text = _tidy(text)

    if not text.strip():
        raise ExtractionError("រកមិនឃើញអត្ថបទនៅក្នុងឯកសារនេះទេ")

    # A document with no Khmer in it is almost certainly the wrong file, and
    # saying so is more useful than filling the editor with English.
    if not has_khmer(text):
        note = note or "ឯកសារនេះហាក់ដូចជាមិនមានអក្សរខ្មែរទេ"

    return text, note


def _classify(mime_type: str, filename: str) -> str:
    mime_type = (mime_type or "").split(";")[0].strip().lower()
    lower_name = filename.lower()

    if mime_type in PDF_TYPES or lower_name.endswith(".pdf"):
        return "pdf"
    if mime_type in DOCX_TYPES or lower_name.endswith(".docx"):
        return "docx"
    if mime_type in TEXT_TYPES or lower_name.endswith((".txt", ".md")):
        return "text"
    if mime_type in IMAGE_TYPES or mime_type.startswith("image/"):
        return "image"
    return "unknown"


def _from_pdf(data: bytes) -> tuple[str, str]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError("ម៉ាស៊ីនមេមិនអាចអានឯកសារ PDF បានទេ") from exc

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        log.warning("could not open PDF: %s", exc)
        raise ExtractionError("មិនអាចបើកឯកសារ PDF នេះបានទេ") from exc

    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as exc:  # noqa: BLE001
            raise ExtractionError("ឯកសារ PDF នេះមានពាក្យសម្ងាត់") from exc

    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:  # noqa: BLE001
            log.debug("page extraction failed: %s", exc)
            pages.append("")

    text = "\n\n".join(p for p in pages if p.strip())

    # A PDF with pages but no extractable text is a scan. The characters are
    # pixels, so this needs the same recognition step images do.
    if not text.strip() and len(reader.pages) > 0:
        raise ExtractionError(
            "ឯកសារ PDF នេះជារូបភាពស្កេន មិនមានអត្ថបទដែលអាចអានបានទេ។"
        )

    return text, ""


def _from_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError("ម៉ាស៊ីនមេមិនអាចអានឯកសារ Word បានទេ") from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        log.warning("could not open docx: %s", exc)
        raise ExtractionError("មិនអាចបើកឯកសារ Word នេះបានទេ") from exc

    parts = [p.text for p in document.paragraphs]

    # Text inside tables is not in document.paragraphs and is easy to lose.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)


def _from_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "utf-16"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ExtractionError("មិនអាចអានអត្ថបទក្នុងឯកសារនេះបានទេ")


def _tidy(text: str) -> str:
    """
    Normalise whitespace without touching the characters themselves.

    Extractors leave ragged spacing and stray blank lines. This only collapses
    runs of whitespace and trims lines; it never rewrites Khmer code points,
    because the editor's underline offsets are computed from this exact string.
    """
    lines = [line.rstrip() for line in text.replace("\r\n", "\n").split("\n")]

    tidied: list[str] = []
    blanks = 0
    for line in lines:
        if line.strip():
            tidied.append(line)
            blanks = 0
        else:
            blanks += 1
            if blanks == 1:
                tidied.append("")

    return "\n".join(tidied).strip()


def decode_payload(encoded: str) -> bytes:
    """Decode a base64 payload, tolerating a data: URL prefix."""
    if "," in encoded and encoded.lstrip().startswith("data:"):
        encoded = encoded.split(",", 1)[1]
    try:
        return base64.b64decode(encoded, validate=False)
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError("ទិន្នន័យឯកសារមិនត្រឹមត្រូវ") from exc
